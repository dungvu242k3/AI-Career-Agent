"""Unit tests for DocxDocumentParser."""

import io
import docx
import pytest
from ai.parsers.docx_parser import (
    DocxDocumentParser,
    DocxParsingError,
    DocxInvalidFormatError,
)


def create_sample_docx_bytes() -> bytes:
    """Helper creating an in-memory sample DOCX with paragraphs, lists, and tables."""
    doc = docx.Document()
    doc.add_heading("Nguyen Van A - Senior AI Engineer", level=1)
    
    # Summary
    doc.add_paragraph("Kỹ sư phần mềm với 4 năm kinh nghiệm xây dựng hệ thống AI và backend phân tán.")
    
    # Skills table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Ngôn ngữ:"
    table.cell(0, 1).text = "Python, Go, TypeScript"
    table.cell(1, 0).text = "Cơ sở dữ liệu:"
    table.cell(1, 1).text = "PostgreSQL, Redis, Qdrant"
    
    # Work Experience
    doc.add_heading("Kinh nghiệm làm việc", level=2)
    doc.add_paragraph("Senior AI Engineer tại VNG Cloud (2023 - Hiện tại)", style="List Bullet")
    doc.add_paragraph("Software Engineer tại MoMo Fintech (2021 - 2023)", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_parser_success():
    """Test extracting clean structured text from valid DOCX bytes."""
    parser = DocxDocumentParser(min_char_count=20)
    docx_bytes = create_sample_docx_bytes()

    text = parser.extract_text_from_bytes(docx_bytes, filename="my_cv.docx")

    assert "Nguyen Van A" in text
    assert "Python, Go, TypeScript" in text
    assert "VNG Cloud" in text
    assert "MoMo Fintech" in text


def test_docx_parser_invalid_magic_bytes():
    """Test rejecting non-docx file (invalid magic bytes)."""
    parser = DocxDocumentParser()
    fake_bytes = b"This is plain text not docx"

    with pytest.raises(DocxInvalidFormatError) as exc_info:
        parser.extract_text_from_bytes(fake_bytes, filename="bad.docx")

    assert "Magic bytes không khớp" in str(exc_info.value)


def test_docx_parser_empty_document():
    """Test rejecting empty DOCX document without content."""
    parser = DocxDocumentParser(min_char_count=50)
    empty_doc = docx.Document()
    buf = io.BytesIO()
    empty_doc.save(buf)

    with pytest.raises(DocxParsingError) as exc_info:
        parser.extract_text_from_bytes(buf.getvalue(), filename="empty.docx")

    assert "không chứa đủ văn bản" in str(exc_info.value)
