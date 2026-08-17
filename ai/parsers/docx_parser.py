"""Word Document (.docx) Parser — High-fidelity text extraction.

Implements BaseDocumentParser using python-docx with paragraph, heading,
bullet-point, and table structure preservation.
"""

import io
import logging
from pathlib import Path
import re
import unicodedata
from docx import Document

from ai.interfaces.parser import BaseDocumentParser

logger = logging.getLogger(__name__)

# Standard ZIP / DOCX magic bytes
DOCX_MAGIC_BYTES = b"\x50\x4B\x03\x04"


class DocxParsingError(Exception):
    """Raised when Word document cannot be parsed."""
    pass


class DocxInvalidFormatError(ValueError):
    """Raised when uploaded file is not a valid DOCX file."""
    pass


class DocxDocumentParser(BaseDocumentParser):
    """Production DOCX document parser for extracting CV text."""

    def __init__(self, min_char_count: int = 40):
        self.min_char_count = min_char_count

    def _normalize_text(self, text: str) -> str:
        """Apply Unicode NFC normalization and clean excessive spaces."""
        text = unicodedata.normalize("NFC", text)
        # Normalize bullet markers
        text = re.sub(r"^[•·▪■►◆-]\s*", "- ", text, flags=re.MULTILINE)
        # Collapse multiple empty lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def extract_text_from_bytes(self, content_bytes: bytes, filename: str = "upload.docx") -> str:
        """Extract text from Word document (.docx) bytes."""
        if not content_bytes or len(content_bytes) < 4:
            raise DocxInvalidFormatError(f"Tệp '{filename}' bị rỗng hoặc không hợp lệ.")

        # Check DOCX ZIP magic bytes
        if not content_bytes.startswith(DOCX_MAGIC_BYTES):
            raise DocxInvalidFormatError(
                f"Tệp '{filename}' không phải là tệp Microsoft Word (.docx) chuẩn. Magic bytes không khớp."
            )

        try:
            doc = Document(io.BytesIO(content_bytes))
        except Exception as e:
            logger.error("Failed to open docx file %s: %s", filename, e)
            raise DocxParsingError(f"Không thể giải mã cấu trúc tệp Word '{filename}': {e}")

        extracted_sections: list[str] = []

        # 1. Extract Paragraphs
        for paragraph in doc.paragraphs:
            p_text = paragraph.text.strip()
            if not p_text:
                continue

            style_name = paragraph.style.name.lower() if paragraph.style else ""
            if "heading 1" in style_name:
                extracted_sections.append(f"\n# {p_text}\n")
            elif "heading 2" in style_name:
                extracted_sections.append(f"\n## {p_text}\n")
            elif "heading 3" in style_name:
                extracted_sections.append(f"\n### {p_text}\n")
            elif "list" in style_name or "bullet" in style_name:
                extracted_sections.append(f"- {p_text}")
            else:
                extracted_sections.append(p_text)

        # 2. Extract Tables (many CVs use tables for Skills or Work History)
        for table in doc.tables:
            table_lines: list[str] = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    # Deduplicate adjacent duplicate cells from merged table cells
                    deduped_cells = []
                    for c in row_cells:
                        if not deduped_cells or deduped_cells[-1] != c:
                            deduped_cells.append(c)
                    table_lines.append(" | ".join(deduped_cells))
            if table_lines:
                extracted_sections.append("\n" + "\n".join(table_lines) + "\n")

        full_text = "\n".join(extracted_sections)
        cleaned_text = self._normalize_text(full_text)

        if len(cleaned_text) < self.min_char_count:
            raise DocxParsingError(
                f"Tệp Word '{filename}' không chứa đủ văn bản có thể đọc được (tìm thấy {len(cleaned_text)} ký tự)."
            )

        return cleaned_text

    def extract_text_from_file(self, file_path: str | Path) -> str:
        """Extract text from a local .docx file."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
        content = path.read_bytes()
        return self.extract_text_from_bytes(content, filename=path.name)
